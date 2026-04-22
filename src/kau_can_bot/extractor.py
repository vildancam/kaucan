from __future__ import annotations

from io import BytesIO
from urllib.parse import urlparse

from bs4 import BeautifulSoup, FeatureNotFound
from pypdf import PdfReader

from .models import PageDocument
from .utils import clean_text, normalize_url


TEXT_TAGS = ("h1", "h2", "h3", "h4", "p", "li", "tr", "td", "th")
CONTAINER_TAGS = ("span", "div")
REMOVE_TAGS = (
    "script",
    "style",
    "noscript",
    "svg",
    "canvas",
    "iframe",
)
NAVIGATION_HINTS = (
    "breadcrumb",
    "dropdown",
    "footer",
    "menu",
    "modal",
    "nav",
    "navbar",
    "pagination",
    "popup",
    "sidebar",
)


def extract_html(html: str, url: str) -> PageDocument:
    try:
        soup = BeautifulSoup(html, "lxml")
    except FeatureNotFound:
        soup = BeautifulSoup(html, "html.parser")

    page_url = normalize_url(url) or url
    links: list[str] = []
    anchor_blocks: list[str] = []
    seen_links: set[str] = set()
    seen_anchor_blocks: set[str] = set()
    for anchor in soup.find_all("a", href=True):
        normalized = normalize_url(anchor["href"], base_url=url)
        if normalized and normalized not in seen_links:
            links.append(normalized)
            seen_links.add(normalized)
        anchor_text = clean_text(anchor.get_text(" ", strip=True))
        anchor_block = _anchor_block(anchor_text, normalized, page_url)
        if anchor_block and anchor_block not in seen_anchor_blocks:
            anchor_blocks.append(anchor_block)
            seen_anchor_blocks.add(anchor_block)

    for tag in soup.find_all(REMOVE_TAGS):
        tag.decompose()

    html_title = clean_text(soup.title.get_text(" ", strip=True)) if soup.title else ""
    heading_title = ""
    for heading in soup.find_all(("h1", "h2", "h3", "h4")):
        if _is_navigation_like(heading):
            continue
        heading_text = clean_text(heading.get_text(" ", strip=True))
        if _is_meaningful(heading_text):
            heading_title = heading_text
            break
    title = heading_title or html_title
    blocks: list[str] = []
    seen_blocks: set[str] = set()

    for tag in soup.find_all(TEXT_TAGS):
        if _is_navigation_like(tag):
            continue
        text = clean_text(tag.get_text(" ", strip=True))
        if _is_meaningful(text) and text not in seen_blocks:
            blocks.append(text)
            seen_blocks.add(text)

    for tag in soup.find_all(CONTAINER_TAGS):
        if _is_navigation_like(tag):
            continue
        if tag.find(TEXT_TAGS):
            continue
        text = clean_text(tag.get_text(" ", strip=True))
        if _is_meaningful(text) and text not in seen_blocks:
            blocks.append(text)
            seen_blocks.add(text)

    for text in anchor_blocks:
        if text not in seen_blocks:
            blocks.append(text)
            seen_blocks.add(text)

    return PageDocument(
        url=url,
        title=title or _fallback_title(url),
        content="\n".join(blocks),
        content_type="text/html",
        links=links,
        metadata={"block_count": len(blocks)},
    )


def extract_pdf(content: bytes, url: str) -> PageDocument:
    title = _fallback_title(url)
    pages: list[str] = []
    metadata: dict[str, str | int] = {}

    try:
        reader = PdfReader(BytesIO(content))
        metadata["page_count"] = len(reader.pages)
        if reader.metadata and reader.metadata.title:
            title = clean_text(reader.metadata.title)

        for page_number, page in enumerate(reader.pages, start=1):
            text = clean_text(page.extract_text() or "")
            if text:
                pages.append(f"Sayfa {page_number}: {text}")
    except Exception as exc:
        metadata["extract_error"] = str(exc)

    return PageDocument(
        url=url,
        title=title,
        content="\n".join(pages),
        content_type="application/pdf",
        links=[],
        metadata=metadata,
    )


def extract_spreadsheet(content: bytes, url: str) -> PageDocument:
    extension = urlparse(url).path.lower().rsplit(".", 1)[-1]
    title = _fallback_title(url)
    rows: list[str] = []
    metadata: dict[str, str | int] = {"document_type": extension}

    try:
        if extension == "xlsx":
            from openpyxl import load_workbook

            workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
            metadata["sheet_count"] = len(workbook.worksheets)
            for sheet in workbook.worksheets:
                rows.append(f"Çalışma sayfası: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    values = [clean_text(str(value)) for value in row if value is not None]
                    if values:
                        rows.append(" | ".join(values))
        elif extension == "xls":
            import xlrd

            workbook = xlrd.open_workbook(file_contents=content)
            metadata["sheet_count"] = workbook.nsheets
            for sheet in workbook.sheets():
                rows.append(f"Çalışma sayfası: {sheet.name}")
                for row_index in range(sheet.nrows):
                    values = [
                        clean_text(str(sheet.cell_value(row_index, col_index)))
                        for col_index in range(sheet.ncols)
                        if sheet.cell_value(row_index, col_index) not in ("", None)
                    ]
                    if values:
                        rows.append(" | ".join(values))
    except Exception as exc:
        metadata["extract_error"] = str(exc)

    return PageDocument(
        url=url,
        title=title,
        content="\n".join(rows),
        content_type=f"application/{extension}",
        links=[],
        metadata=metadata,
    )


def extract_docx(content: bytes, url: str) -> PageDocument:
    title = _fallback_title(url)
    blocks: list[str] = []
    metadata: dict[str, str] = {"document_type": "docx"}

    try:
        from docx import Document

        document = Document(BytesIO(content))
        for paragraph in document.paragraphs:
            text = clean_text(paragraph.text)
            if text:
                blocks.append(text)
        for table in document.tables:
            for row in table.rows:
                values = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
                if values:
                    blocks.append(" | ".join(values))
    except Exception as exc:
        metadata["extract_error"] = str(exc)

    return PageDocument(
        url=url,
        title=title,
        content="\n".join(blocks),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        links=[],
        metadata=metadata,
    )


def extract_plain_text(content: bytes, url: str, content_type: str) -> PageDocument:
    text = content.decode("utf-8", errors="ignore")
    return PageDocument(
        url=url,
        title=_fallback_title(url),
        content=clean_text(text),
        content_type=content_type,
        links=[],
        metadata={},
    )


def unsupported_document(url: str, content_type: str) -> PageDocument:
    return PageDocument(
        url=url,
        title=_fallback_title(url),
        content="",
        content_type=content_type,
        links=[],
        metadata={"note": "Bu ek türünden metin çıkarımı desteklenmiyor."},
    )


def _is_meaningful(text: str) -> bool:
    if len(text) < 3:
        return False
    if len(text) > 1400:
        return False
    if text.count(" ") == 0 and len(text) < 12:
        return False
    return True


def _anchor_block(text: str, url: str | None, page_url: str) -> str | None:
    if not url:
        return None
    if url == page_url:
        return None
    text = clean_text(text)
    if not _is_meaningful(text) or len(text) > 140:
        return None
    normalized = text.lower()
    if normalized in {
        "devamını oku",
        "duyuruyu paylaş:",
        "facebook",
        "twitter",
        "yeni pencerede aç",
    }:
        return None
    path = urlparse(url).path.lower()
    if not (
        "/iibf" in path
        or path.startswith("/dersler.aspx")
        or path.startswith("/kau/rehber")
    ):
        return None
    return f"Bağlantı: {text} | URL: {url}"


def _is_navigation_like(tag) -> bool:
    for parent in [tag, *tag.parents]:
        if getattr(parent, "name", None) == "nav":
            return True

        attrs = " ".join(
            str(value)
            for key in ("id", "class", "role", "aria-label")
            for value in (
                parent.get(key, [])
                if isinstance(parent.get(key, []), list)
                else [parent.get(key, "")]
            )
        ).lower()
        if any(hint in attrs for hint in NAVIGATION_HINTS):
            return True

    return False


def _fallback_title(url: str) -> str:
    path = urlparse(url).path.rstrip("/").split("/")[-1]
    return clean_text(path.replace("-", " ").replace("_", " ")) or url
